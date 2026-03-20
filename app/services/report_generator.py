import io
import json
import logging
import os
from datetime import datetime, timezone
from typing import Any

from openai import AsyncOpenAI
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from app.config import get_settings
from app.models.session import Session, Transcript, Rubric, Analysis
from app.models.report import Report

logger = logging.getLogger(__name__)
settings = get_settings()

_openai_client: AsyncOpenAI | None = None


def _get_openai() -> AsyncOpenAI:
    global _openai_client
    if _openai_client is None:
        _openai_client = AsyncOpenAI(api_key=settings.OPENAI_API_KEY)
    return _openai_client


async def _get_coaching_content(analysis: dict, rubric: dict) -> dict:
    """
    Use GPT-4o to generate coaching content in both English and Spanish:
    - overall_assessment: 2-3 sentence narrative summary of the agent's performance
    - recommendations: 3-5 specific, actionable coaching recommendations
    """
    client = _get_openai()

    system_prompt = (
        "You are an expert call center coach. "
        "Based on the analysis results, write a brief overall assessment paragraph and "
        "3-5 specific, actionable coaching recommendations. "
        "Return BOTH in English and Spanish. Return JSON only."
    )
    user_prompt = f"""
Analysis:
{json.dumps(analysis, indent=2)}

Rubric criteria:
{json.dumps(rubric.get("criteria", []), indent=2)}

Return a JSON object with this exact structure:
{{
  "overall_assessment": {{
    "en": "2-3 sentence narrative summary of the agent's overall performance in English.",
    "es": "Resumen narrativo de 2-3 oraciones del desempeño general del agente en español."
  }},
  "recommendations": {{
    "en": ["Actionable recommendation 1", "Actionable recommendation 2"],
    "es": ["Recomendación accionable 1", "Recomendación accionable 2"]
  }}
}}
"""
    response = await client.chat.completions.create(
        model="gpt-4o",
        response_format={"type": "json_object"},
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.3,
        max_tokens=2048,
    )
    parsed = json.loads(response.choices[0].message.content)
    assessment = parsed.get("overall_assessment", {})
    recs = parsed.get("recommendations", {})
    return {
        "overall_assessment": {
            "en": str(assessment.get("en", "")),
            "es": str(assessment.get("es", "")),
        },
        "recommendations": {
            "en": [str(r) for r in recs.get("en", [])],
            "es": [str(r) for r in recs.get("es", [])],
        },
    }


async def generate_report(session_id: str, db: AsyncSession) -> dict:
    """
    Compile a full bilingual coaching report from all stored data for the session.
    Returns a report dict and upserts it into the DB.
    """
    session_result = await db.execute(select(Session).where(Session.id == session_id))
    session: Session = session_result.scalar_one_or_none()
    if not session:
        raise ValueError(f"Session {session_id} not found")

    transcript_result = await db.execute(select(Transcript).where(Transcript.session_id == session_id))
    transcript_row: Transcript | None = transcript_result.scalar_one_or_none()
    utterances = transcript_row.utterances if transcript_row else []

    rubric_result = await db.execute(select(Rubric).where(Rubric.session_id == session_id))
    rubric_row: Rubric | None = rubric_result.scalar_one_or_none()
    rubric = rubric_row.criteria if rubric_row else {"criteria": []}

    analysis_result = await db.execute(select(Analysis).where(Analysis.session_id == session_id))
    analysis_row: Analysis | None = analysis_result.scalar_one_or_none()
    if not analysis_row:
        raise ValueError(f"Analysis for session {session_id} not found")

    scores = analysis_row.scores or []
    strengths = analysis_row.strengths or {"en": [], "es": []}
    improvements = analysis_row.improvements or {"en": [], "es": []}
    key_moments = analysis_row.key_moments or []

    # Normalise strengths/improvements to dict form (handle old flat-array rows)
    if isinstance(strengths, list):
        strengths = {"en": strengths, "es": []}
    if isinstance(improvements, list):
        improvements = {"en": improvements, "es": []}

    analysis_data = {
        "scores": scores,
        "strengths": strengths,
        "improvements": improvements,
        "key_moments": key_moments,
    }

    total_score = sum(s.get("score", 0) for s in scores)
    total_max = sum(s.get("max_score", 10) for s in scores)

    coaching = await _get_coaching_content(analysis_data, rubric)
    recommendations = coaching["recommendations"]
    overall_assessment = coaching["overall_assessment"]

    meta_es = getattr(session, "metadata_es", None) or {}

    report_data: dict[str, Any] = {
        "summary": {
            "agent_name": session.agent_name or "Unknown Agent",
            "agent_name_es": meta_es.get("agent_name") or session.agent_name or "Agente Desconocido",
            "client_name": session.client_name or "Unknown Client",
            "client_name_es": meta_es.get("client_name") or session.client_name or "Cliente Desconocido",
            "call_title": session.call_title or "Untitled Call",
            "call_title_es": meta_es.get("call_title") or session.call_title or "Llamada sin título",
            "call_date": session.call_date or datetime.now(timezone.utc).strftime("%Y-%m-%d"),
            "overall_score": round(total_score, 2),
            "max_score": total_max,
            "percentage": round((total_score / total_max * 100) if total_max > 0 else 0, 1),
        },
        "overall_assessment": overall_assessment,  # {"en": "...", "es": "..."}
        "rubric_scores": scores,          # each item has category, category_es, reason, reason_es
        "strengths": strengths,           # {"en": [...], "es": [...]}
        "areas_for_improvement": improvements,  # {"en": [...], "es": [...]}
        "key_moments": key_moments,       # each item has description, description_es
        "recommendations": recommendations,     # {"en": [...], "es": [...]}
        "transcript_summary": {
            "total_utterances": len(utterances),
            "agent_turns": sum(1 for u in utterances if u.get("speaker") == "Agent"),
            "customer_turns": sum(1 for u in utterances if u.get("speaker") == "Customer"),
        },
    }

    report_result = await db.execute(select(Report).where(Report.session_id == session_id))
    existing_report: Report | None = report_result.scalar_one_or_none()
    if existing_report:
        existing_report.report_data = report_data
        db.add(existing_report)
    else:
        new_report = Report(session_id=session_id, report_data=report_data)
        db.add(new_report)

    await db.commit()
    logger.info(f"Bilingual report generated and saved for session {session_id}")
    return report_data


# ---------------------------------------------------------------------------
# Export helpers
# ---------------------------------------------------------------------------

def _pick(bilingual: Any, lang: str) -> list:
    """Return the list for the given language from a {"en": [...], "es": [...]} dict, or the value as-is."""
    if isinstance(bilingual, dict):
        return bilingual.get(lang) or bilingual.get("en") or []
    return bilingual or []


def generate_pdf_bytes(report_data: dict, lang: str = "en") -> bytes:
    """Generate a PDF report from report_data and return as bytes."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
        HRFlowable,
        KeepTogether,
    )

    is_es = lang == "es"
    PAGE_W, PAGE_H = A4
    LEFT_MARGIN = 2 * cm
    RIGHT_MARGIN = 2 * cm
    TOP_MARGIN = 2 * cm
    BOTTOM_MARGIN = 2 * cm
    USABLE_WIDTH = PAGE_W - LEFT_MARGIN - RIGHT_MARGIN

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=LEFT_MARGIN,
        rightMargin=RIGHT_MARGIN,
        topMargin=TOP_MARGIN,
        bottomMargin=BOTTOM_MARGIN,
    )
    styles = getSampleStyleSheet()

    heading1 = ParagraphStyle("Heading1Custom", parent=styles["Heading1"], fontSize=18, spaceAfter=6, spaceBefore=0, leading=22, wordWrap="CJK")
    heading2 = ParagraphStyle("Heading2Custom", parent=styles["Heading2"], fontSize=13, spaceAfter=4, spaceBefore=10, leading=18, textColor=colors.HexColor("#C0392B"), wordWrap="CJK")
    normal = ParagraphStyle("NormalCustom", parent=styles["Normal"], fontSize=9, leading=14, spaceAfter=3, wordWrap="CJK")
    bullet = ParagraphStyle("BulletCustom", parent=normal, leftIndent=14, firstLineIndent=0, spaceAfter=4)
    label_style = ParagraphStyle("LabelStyle", parent=normal, fontName="Helvetica-Bold", fontSize=9, leading=14)

    story = []

    title_text = "Informe de Coaching — CallMentor AI" if is_es else "CallMentor AI — Coaching Report"
    story.append(Paragraph(title_text, heading1))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#C0392B")))
    story.append(Spacer(1, 0.4 * cm))

    summary = report_data.get("summary", {})
    story.append(Paragraph("Resumen de la llamada" if is_es else "Call Summary", heading2))

    label_w = 3.8 * cm
    value_w = USABLE_WIDTH - label_w

    def _label(text: str) -> Paragraph:
        return Paragraph(text, label_style)

    def _value(text: str) -> Paragraph:
        return Paragraph(str(text), normal)

    if is_es:
        agent_label, client_label, title_label, date_label, score_label = "Agente", "Cliente", "Título", "Fecha", "Puntuación Total"
        agent_val = summary.get("agent_name_es") or summary.get("agent_name", "—")
        client_val = summary.get("client_name_es") or summary.get("client_name", "—")
        title_val = summary.get("call_title_es") or summary.get("call_title", "—")
    else:
        agent_label, client_label, title_label, date_label, score_label = "Agent", "Client", "Call Title", "Call Date", "Overall Score"
        agent_val = summary.get("agent_name", "—")
        client_val = summary.get("client_name", "—")
        title_val = summary.get("call_title", "—")

    summary_rows = [
        [_label(agent_label),  _value(agent_val)],
        [_label(client_label), _value(client_val)],
        [_label(title_label),  _value(title_val)],
        [_label(date_label),   _value(summary.get("call_date", "—"))],
        [_label(score_label),  _value(f"{summary.get('overall_score', 0)} / {summary.get('max_score', 0)} ({summary.get('percentage', 0)}%)")],
    ]
    tbl = Table(summary_rows, colWidths=[label_w, value_w])
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F2F2F2")),
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTNAME", (1, 0), (1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CCCCCC")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 7),
        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
    ]))
    story.append(tbl)
    story.append(Spacer(1, 0.6 * cm))

    scores = report_data.get("rubric_scores", [])
    if scores:
        story.append(Paragraph("Puntuaciones de Evaluación" if is_es else "Evaluation Scores", heading2))
        cat_w, score_w, max_w = 4.0 * cm, 1.6 * cm, 1.6 * cm
        reason_w = USABLE_WIDTH - cat_w - score_w - max_w

        header_style = ParagraphStyle("HeaderStyle", parent=normal, fontName="Helvetica-Bold", textColor=colors.white, fontSize=9)
        if is_es:
            col_headers = ["Categoría", "Punt.", "Máx.", "Razón"]
        else:
            col_headers = ["Category", "Score", "Max", "Reason"]

        score_rows = [[Paragraph(h, header_style) for h in col_headers]]
        for s in scores:
            cat = s.get("category_es") or s.get("category", "") if is_es else s.get("category", "")
            reason = s.get("reason_es") or s.get("reason", "") if is_es else s.get("reason", "")
            score_rows.append([
                Paragraph(cat, normal),
                Paragraph(str(s.get("score", 0)), normal),
                Paragraph(str(s.get("max_score", 0)), normal),
                Paragraph(reason, normal),
            ])

        score_tbl = Table(score_rows, colWidths=[cat_w, score_w, max_w, reason_w], repeatRows=1)
        score_tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#C0392B")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CCCCCC")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#FFF5F5")]),
        ]))
        story.append(score_tbl)
        story.append(Spacer(1, 0.6 * cm))

    strengths = _pick(report_data.get("strengths", []), lang)
    if strengths:
        items = [Paragraph("Fortalezas" if is_es else "Strengths", heading2)]
        for item in strengths:
            items.append(Paragraph(f"• {item}", bullet))
        items.append(Spacer(1, 0.3 * cm))
        story.append(KeepTogether(items))

    improvements = _pick(report_data.get("areas_for_improvement", []), lang)
    if improvements:
        items = [Paragraph("Áreas de Mejora" if is_es else "Areas for Improvement", heading2)]
        for item in improvements:
            items.append(Paragraph(f"• {item}", bullet))
        items.append(Spacer(1, 0.3 * cm))
        story.append(KeepTogether(items))

    key_moments = report_data.get("key_moments", [])
    if key_moments:
        story.append(Paragraph("Momentos Clave" if is_es else "Key Moments", heading2))
        for km in key_moments:
            ts = km.get("timestamp", "")
            desc = km.get("description_es") or km.get("description", "") if is_es else km.get("description", "")
            text = f"<b>[{ts}]</b> {desc}" if ts else desc
            story.append(Paragraph(text, bullet))
        story.append(Spacer(1, 0.4 * cm))

    recommendations = _pick(report_data.get("recommendations", []), lang)
    if recommendations:
        items = [Paragraph("Recomendaciones de Coaching" if is_es else "Coaching Recommendations", heading2)]
        for i, rec in enumerate(recommendations, 1):
            items.append(Paragraph(f"{i}. {rec}", bullet))
        items.append(Spacer(1, 0.3 * cm))
        story.append(KeepTogether(items))

    doc.build(story)
    buffer.seek(0)
    return buffer.read()


def _safe(value) -> str:
    """Strip XML-incompatible characters from a string."""
    import re
    text = str(value) if value is not None else ""
    return re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)


def generate_docx_bytes(report_data: dict, lang: str = "en") -> bytes:
    """Generate a DOCX report from report_data and return as bytes."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    is_es = lang == "es"
    doc = Document()

    title_text = "Informe de Coaching — CallMentor AI" if is_es else "CallMentor AI — Coaching Report"
    title = doc.add_heading(title_text, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    summary = report_data.get("summary", {})
    doc.add_heading("Resumen de la llamada" if is_es else "Call Summary", level=1)
    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = "Table Grid"

    if is_es:
        agent_val = summary.get("agent_name_es") or summary.get("agent_name", "")
        client_val = summary.get("client_name_es") or summary.get("client_name", "")
        title_val = summary.get("call_title_es") or summary.get("call_title", "")
        rows_data = [
            ("Agente", agent_val),
            ("Cliente", client_val),
            ("Título", title_val),
            ("Fecha", summary.get("call_date", "")),
            ("Puntuación Total", f"{summary.get('overall_score', 0)} / {summary.get('max_score', 0)} ({summary.get('percentage', 0)}%)"),
        ]
    else:
        rows_data = [
            ("Agent", _safe(summary.get("agent_name", ""))),
            ("Client", _safe(summary.get("client_name", ""))),
            ("Call Title", _safe(summary.get("call_title", ""))),
            ("Call Date", _safe(summary.get("call_date", ""))),
            ("Overall Score", f"{summary.get('overall_score', 0)} / {summary.get('max_score', 0)} ({summary.get('percentage', 0)}%)"),
        ]

    for i, (label, value) in enumerate(rows_data):
        tbl.rows[i].cells[0].text = label
        tbl.rows[i].cells[1].text = _safe(value)

    doc.add_paragraph()

    scores = report_data.get("rubric_scores", [])
    if scores:
        doc.add_heading("Puntuaciones de Evaluación" if is_es else "Evaluation Scores", level=1)
        score_tbl = doc.add_table(rows=1 + len(scores), cols=4)
        score_tbl.style = "Table Grid"
        headers = ["Categoría", "Punt.", "Máx.", "Razón"] if is_es else ["Category", "Score", "Max", "Reason"]
        for j, h in enumerate(headers):
            cell = score_tbl.rows[0].cells[j]
            cell.text = h
            run = cell.paragraphs[0].runs[0]
            run.bold = True
        for i, s in enumerate(scores, 1):
            row = score_tbl.rows[i]
            cat = s.get("category_es") or s.get("category", "") if is_es else s.get("category", "")
            reason = s.get("reason_es") or s.get("reason", "") if is_es else s.get("reason", "")
            row.cells[0].text = _safe(cat)
            row.cells[1].text = _safe(s.get("score", 0))
            row.cells[2].text = _safe(s.get("max_score", 0))
            row.cells[3].text = _safe(reason)
        doc.add_paragraph()

    strengths = _pick(report_data.get("strengths", []), lang)
    if strengths:
        doc.add_heading("Fortalezas" if is_es else "Strengths", level=1)
        for item in strengths:
            doc.add_paragraph(_safe(item), style="List Bullet")
        doc.add_paragraph()

    improvements = _pick(report_data.get("areas_for_improvement", []), lang)
    if improvements:
        doc.add_heading("Áreas de Mejora" if is_es else "Areas for Improvement", level=1)
        for item in improvements:
            doc.add_paragraph(_safe(item), style="List Bullet")
        doc.add_paragraph()

    key_moments = report_data.get("key_moments", [])
    if key_moments:
        doc.add_heading("Momentos Clave" if is_es else "Key Moments", level=1)
        for km in key_moments:
            desc = km.get("description_es") or km.get("description", "") if is_es else km.get("description", "")
            doc.add_paragraph(f"[{_safe(km.get('timestamp', ''))}] {_safe(desc)}", style="List Bullet")
        doc.add_paragraph()

    recommendations = _pick(report_data.get("recommendations", []), lang)
    if recommendations:
        doc.add_heading("Recomendaciones de Coaching" if is_es else "Coaching Recommendations", level=1)
        for i, rec in enumerate(recommendations, 1):
            doc.add_paragraph(f"{i}. {_safe(rec)}")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# HTML export
# ---------------------------------------------------------------------------

def _esc(value: Any) -> str:
    """HTML-escape a value for safe inline insertion."""
    import html
    return html.escape(str(value) if value is not None else "")


def _score_color(pct: float) -> str:
    if pct >= 70:
        return "#22c55e"
    if pct >= 45:
        return "#f59e0b"
    return "#ef4444"


def generate_html_bytes(report_data: dict, lang: str = "en") -> bytes:
    """
    Generate a polished, self-contained HTML coaching report inspired by the
    Takeda / Profuturo template styles shared by the team.
    """
    is_es = lang == "es"
    summary = report_data.get("summary", {})

    # ── Localised strings ───────────────────────────────────────────────────
    L: dict[str, str]
    if is_es:
        L = {
            "report_label": "Informe de Evaluación de Coaching",
            "agent": "Agente", "client": "Cliente",
            "date": "Fecha", "score_label": "Calificación global",
            "assessment_title": "Evaluación general del desempeño",
            "strengths": "Fortalezas observadas",
            "improvements": "Áreas de mejora",
            "criteria_title": "Detalle por criterio de evaluación",
            "col_criterion": "Criterio", "col_score": "Puntaje",
            "col_max": "Máx.", "col_reason": "Justificación",
            "moments_title": "Momentos clave",
            "recs_title": "Recomendaciones de coaching",
            "footer": "Este informe fue generado automáticamente por CallMentor AI a partir de la transcripción de la llamada.",
        }
    else:
        L = {
            "report_label": "Coaching Evaluation Report",
            "agent": "Agent", "client": "Client",
            "date": "Date", "score_label": "Overall Score",
            "assessment_title": "Overall Performance Assessment",
            "strengths": "Observed Strengths",
            "improvements": "Areas for Improvement",
            "criteria_title": "Evaluation Criteria Detail",
            "col_criterion": "Criterion", "col_score": "Score",
            "col_max": "Max", "col_reason": "Justification",
            "moments_title": "Key Moments",
            "recs_title": "Coaching Recommendations",
            "footer": "This report was automatically generated by CallMentor AI from the call transcript.",
        }

    # ── Data extraction ─────────────────────────────────────────────────────
    call_title = (summary.get("call_title_es") or summary.get("call_title", "")) if is_es else summary.get("call_title", "")
    agent_name = (summary.get("agent_name_es") or summary.get("agent_name", "")) if is_es else summary.get("agent_name", "")
    client_name = (summary.get("client_name_es") or summary.get("client_name", "")) if is_es else summary.get("client_name", "")
    call_date = summary.get("call_date", "")
    overall = summary.get("overall_score", 0)
    max_score = summary.get("max_score", 1) or 1
    pct = round((overall / max_score) * 100, 1)
    meter_w = min(100, round((overall / max_score) * 100))
    overall_color = _score_color(pct)

    assessment_raw = report_data.get("overall_assessment", {})
    assessment = (assessment_raw.get("es") or assessment_raw.get("en", "")) if is_es else assessment_raw.get("en", "")

    strengths = _pick(report_data.get("strengths", []), lang)
    improvements = _pick(report_data.get("areas_for_improvement", []), lang)
    scores = report_data.get("rubric_scores", [])
    key_moments = report_data.get("key_moments", [])
    recommendations = _pick(report_data.get("recommendations", []), lang)

    # ── Builder helpers ─────────────────────────────────────────────────────
    def bullet_list(items: list[str]) -> str:
        if not items:
            return ""
        return "<ul>" + "".join(f"<li>{_esc(i)}</li>" for i in items) + "</ul>"

    def score_rows_html() -> str:
        rows = []
        for s in scores:
            cat = (s.get("category_es") or s.get("category", "")) if is_es else s.get("category", "")
            reason = (s.get("reason_es") or s.get("reason", "")) if is_es else s.get("reason", "")
            sc = s.get("score", 0)
            mx = s.get("max_score", 10) or 10
            sp = round((sc / mx) * 100)
            badge_color = _score_color(sp)
            rows.append(f"""
              <tr>
                <td class="col-crit">{_esc(cat)}</td>
                <td class="col-score">
                  <span class="badge" style="background:{badge_color}18;color:{badge_color};border-color:{badge_color}44">
                    {_esc(sc)}&thinsp;/&thinsp;{_esc(mx)}
                  </span>
                </td>
                <td>{_esc(reason)}</td>
              </tr>""")
        return "".join(rows)

    def moments_html() -> str:
        if not key_moments:
            return ""
        items = []
        for km in key_moments:
            ts = km.get("timestamp", "")
            desc = (km.get("description_es") or km.get("description", "")) if is_es else km.get("description", "")
            ts_html = f'<span class="ts-badge">{_esc(ts)}</span>' if ts else ""
            items.append(f'<li>{ts_html}<span class="km-desc">{_esc(desc)}</span></li>')
        return "<ul class='km-list'>" + "".join(items) + "</ul>"

    def recs_html() -> str:
        return "<ol>" + "".join(f"<li>{_esc(r)}</li>" for r in recommendations) + "</ol>"

    # ── Pre-compute conditional HTML blocks (avoid nested f-string escaping) ─
    strengths_block = ""
    if strengths:
        strengths_block = f"""      <div class="card strengths-card">
        <div class="card-title">
          <svg width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><polyline points="20 6 9 17 4 12"/></svg>
          {_esc(L["strengths"])}
        </div>
        {bullet_list(strengths)}
      </div>"""

    improvements_block = ""
    if improvements:
        improvements_block = f"""      <div class="card improvements-card">
        <div class="card-title">
          <svg width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><circle cx="12" cy="12" r="10"/><line x1="12" y1="8" x2="12" y2="12"/><line x1="12" y1="16" x2="12.01" y2="16"/></svg>
          {_esc(L["improvements"])}
        </div>
        {bullet_list(improvements)}
      </div>"""

    two_col_block = ""
    if strengths or improvements:
        two_col_block = f"""    <div class="two-col">
{strengths_block}
{improvements_block}
    </div>"""

    # ── Meta info chips ─────────────────────────────────────────────────────
    meta_chips = ""
    if agent_name:
        meta_chips += f'<span class="chip"><svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><circle cx="12" cy="8" r="5"/><path d="M20 21a8 8 0 1 0-16 0"/></svg>{_esc(L["agent"])}: {_esc(agent_name)}</span>'
    if client_name:
        meta_chips += f'<span class="chip"><svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><circle cx="12" cy="8" r="5"/><path d="M20 21a8 8 0 1 0-16 0"/></svg>{_esc(L["client"])}: {_esc(client_name)}</span>'
    if call_date:
        meta_chips += f'<span class="chip"><svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><rect x="3" y="4" width="18" height="18" rx="2"/><line x1="16" y1="2" x2="16" y2="6"/><line x1="8" y1="2" x2="8" y2="6"/><line x1="3" y1="10" x2="21" y2="10"/></svg>{_esc(call_date)}</span>'

    # ── Full HTML ───────────────────────────────────────────────────────────
    html = f"""<!DOCTYPE html>
<html lang="{lang}">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{_esc(call_title or L["report_label"])}</title>
<style>
  *, *::before, *::after {{ box-sizing: border-box; margin: 0; padding: 0; }}
  body {{ font-family: Arial, Helvetica, sans-serif; background: #f4f4f6; color: #2d2d2d; line-height: 1.5; }}
  .shell {{ max-width: 1100px; margin: 32px auto; background: #fff; border-radius: 20px; overflow: hidden; box-shadow: 0 12px 40px rgba(0,0,0,0.10); border: 1px solid #ddd; }}

  /* ── Header ── */
  .header {{ background: linear-gradient(135deg, #c0392b 0%, #96161c 100%); color: #fff; padding: 28px 32px 24px; position: relative; overflow: hidden; }}
  .header::after {{ content:""; position:absolute; right:-60px; top:-60px; width:220px; height:220px; border-radius:50%; background:rgba(255,255,255,0.07); pointer-events:none; }}
  .header-inner {{ position: relative; z-index: 1; }}
  .kicker {{ font-size: 11px; font-weight: 700; letter-spacing: .12em; text-transform: uppercase; color: rgba(255,255,255,.8); margin-bottom: 6px; }}
  .header h1 {{ font-size: 26px; font-weight: 900; line-height: 1.15; }}
  .chips {{ display: flex; flex-wrap: wrap; gap: 10px; margin-top: 12px; }}
  .chip {{ display: inline-flex; align-items: center; gap: 5px; font-size: 13px; color: rgba(255,255,255,.9); background: rgba(255,255,255,.15); border-radius: 999px; padding: 4px 12px; }}

  /* ── Hero ── */
  .hero {{ display: grid; grid-template-columns: 260px 1fr; gap: 20px; padding: 24px 32px; background: linear-gradient(180deg,#fff 0%,#f9f9f9 100%); border-bottom: 1px solid #eee; }}
  .score-card {{ background: linear-gradient(180deg,#fff 0%,#fdf1f1 100%); border: 1px solid #e8c8c8; border-radius: 16px; padding: 20px; text-align: center; box-shadow: 0 4px 16px rgba(192,57,43,.08); }}
  .score-label {{ font-size: 11px; font-weight: 700; letter-spacing: .08em; text-transform: uppercase; color: #888; margin-bottom: 8px; }}
  .score-num {{ font-size: 64px; font-weight: 900; line-height: 1; color: {overall_color}; }}
  .score-denom {{ font-size: 18px; font-weight: 700; color: #aaa; }}
  .meter {{ margin-top: 14px; height: 10px; border-radius: 999px; background: #ececec; overflow: hidden; border: 1px solid #ddd; }}
  .meter-fill {{ height: 100%; width: {meter_w}%; background: linear-gradient(90deg, #f5a8a8 0%, {overall_color} 60%, #8a1a12 100%); border-radius: 999px; }}
  .pct-label {{ margin-top: 8px; font-size: 13px; font-weight: 700; color: {overall_color}; }}

  .assessment-card {{ background: #f8f8f8; border: 1px solid #e5e5e5; border-radius: 16px; padding: 22px; display: flex; flex-direction: column; justify-content: center; }}
  .assessment-card h3 {{ font-size: 16px; font-weight: 800; color: #3a3a3a; margin-bottom: 10px; }}
  .assessment-card p {{ font-size: 14px; line-height: 1.7; color: #555; white-space: pre-wrap; }}

  /* ── Content ── */
  .content {{ padding: 28px 32px; background: #f7f7f8; display: flex; flex-direction: column; gap: 22px; }}

  .card {{ background: #fff; border: 1px solid #e5e5e5; border-radius: 16px; padding: 22px; box-shadow: 0 3px 12px rgba(0,0,0,.04); }}
  .card-title {{ font-size: 15px; font-weight: 800; color: #2d2d2d; margin-bottom: 14px; padding-bottom: 10px; border-bottom: 1px solid #eee; display: flex; align-items: center; gap: 8px; }}
  .card-title svg {{ flex-shrink: 0; }}

  .two-col {{ display: grid; grid-template-columns: 1fr 1fr; gap: 18px; }}

  .strengths-card .card-title {{ color: #16a34a; }}
  .improvements-card .card-title {{ color: #d97706; }}

  ul {{ padding-left: 18px; }}
  ul li {{ font-size: 14px; line-height: 1.6; color: #444; margin-bottom: 6px; }}
  ol {{ padding-left: 18px; }}
  ol li {{ font-size: 14px; line-height: 1.6; color: #444; margin-bottom: 8px; }}

  /* ── Criteria table ── */
  .table-wrap {{ overflow-x: auto; }}
  table {{ width: 100%; border-collapse: collapse; min-width: 600px; }}
  thead th {{ background: linear-gradient(180deg,#f5f5f5 0%,#ebebeb 100%); font-size: 12px; font-weight: 700; text-transform: uppercase; letter-spacing: .06em; color: #555; padding: 11px 12px; border: 1px solid #ddd; text-align: left; }}
  tbody td {{ padding: 12px; border: 1px solid #e8e8e8; font-size: 14px; line-height: 1.5; vertical-align: top; color: #333; }}
  tbody tr:nth-child(even) td {{ background: #fafafa; }}
  .col-crit {{ width: 25%; font-weight: 600; }}
  .col-score {{ width: 12%; text-align: center; }}
  .badge {{ display: inline-flex; align-items: center; justify-content: center; padding: 6px 10px; border-radius: 999px; font-size: 13px; font-weight: 800; border: 1px solid; white-space: nowrap; }}

  /* ── Key moments ── */
  .km-list {{ list-style: none; padding: 0; display: flex; flex-direction: column; gap: 10px; }}
  .km-list li {{ display: flex; align-items: flex-start; gap: 10px; font-size: 14px; color: #444; }}
  .ts-badge {{ flex-shrink: 0; font-family: monospace; font-size: 12px; font-weight: 700; background: #fdf1f1; color: #c0392b; border: 1px solid #f0c8c8; border-radius: 6px; padding: 2px 8px; margin-top: 2px; }}
  .km-desc {{ line-height: 1.6; }}

  /* ── Footer ── */
  .footer {{ padding: 16px 32px; border-top: 1px solid #eee; background: #f9f9f9; font-size: 12px; color: #999; text-align: center; }}

  @media (max-width: 780px) {{
    .hero, .two-col {{ grid-template-columns: 1fr; }}
    .shell {{ margin: 0; border-radius: 0; }}
    .content, .header {{ padding: 18px; }}
    .hero {{ padding: 18px; }}
    .score-num {{ font-size: 48px; }}
  }}
</style>
</head>
<body>
<div class="shell">

  <div class="header">
    <div class="header-inner">
      <div class="kicker">{_esc(L["report_label"])}</div>
      <h1>{_esc(call_title or L["report_label"])}</h1>
      <div class="chips">{meta_chips}</div>
    </div>
  </div>

  <div class="hero">
    <div class="score-card">
      <div class="score-label">{_esc(L["score_label"])}</div>
      <div class="score-num">{_esc(overall)}<span class="score-denom">&thinsp;/&thinsp;{_esc(max_score)}</span></div>
      <div class="meter"><div class="meter-fill"></div></div>
      <div class="pct-label">{pct}%</div>
    </div>
    <div class="assessment-card">
      <h3>{_esc(L["assessment_title"])}</h3>
      <p>{_esc(assessment) if assessment else "—"}</p>
    </div>
  </div>

  <div class="content">

    {two_col_block}

    {"" if not scores else f'''
    <div class="card">
      <div class="card-title">
        <svg width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><rect x="3" y="3" width="18" height="18" rx="2"/><line x1="3" y1="9" x2="21" y2="9"/><line x1="9" y1="21" x2="9" y2="9"/></svg>
        {_esc(L["criteria_title"])}
      </div>
      <div class="table-wrap">
        <table>
          <thead>
            <tr>
              <th class="col-crit">{_esc(L["col_criterion"])}</th>
              <th class="col-score">{_esc(L["col_score"])}&nbsp;/&nbsp;{_esc(L["col_max"])}</th>
              <th>{_esc(L["col_reason"])}</th>
            </tr>
          </thead>
          <tbody>{score_rows_html()}</tbody>
        </table>
      </div>
    </div>'''}

    {"" if not key_moments else f'''
    <div class="card">
      <div class="card-title">
        <svg width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><circle cx="12" cy="12" r="10"/><polyline points="12 6 12 12 16 14"/></svg>
        {_esc(L["moments_title"])}
      </div>
      {moments_html()}
    </div>'''}

    {"" if not recommendations else f'''
    <div class="card">
      <div class="card-title">
        <svg width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><path d="M9 11l3 3L22 4"/><path d="M21 12v7a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2V5a2 2 0 0 1 2-2h11"/></svg>
        {_esc(L["recs_title"])}
      </div>
      {recs_html()}
    </div>'''}

  </div>

  <div class="footer">{_esc(L["footer"])}</div>

</div>
</body>
</html>"""

    return html.encode("utf-8")
